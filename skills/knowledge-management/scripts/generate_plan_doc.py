#!/usr/bin/env python3
"""将 knowledge-system-plan.md 转换为 Word 文档（组织汇报版）"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ========== 页面设置：横向 A4 ==========
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ========== 辅助函数 ==========
def set_run_font(run, name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    for child in rPr:
        if child.tag.endswith('}rFonts'):
            child.set(qn('w:eastAsia'), name)

def set_cell_bg(cell, fill):
    """设置表格单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 移除已有的 shd
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    return p

def add_para(doc, text, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_table_header(table, headers, bg='003366'):
    """添加表头行"""
    hdr = table.rows[0]
    for cell, txt in zip(hdr.cells, headers):
        cell.text = txt
        set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(txt)
        set_run_font(run, size=9, bold=True, color=(255, 255, 255))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_data_row(table, values, font_size=9):
    """添加数据行"""
    row = table.add_row()
    for cell, txt in zip(row.cells, values):
        cell.text = txt
        set_run_font(cell.paragraphs[0].runs[0], size=font_size)

# ========== 封面 ==========
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('知识管理体系规划方案')
set_run_font(run, size=28, bold=True, color=(0, 51, 102))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 以 OpenClaw 为核心的能源数据要素全链路知识中枢')
set_run_font(run, size=14, color=(80, 80, 80))

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('编制人：Domi（AI 助手）\n编制日期：2026年4月21日')
set_run_font(run, size=11, color=(100, 100, 100))

doc.add_page_break()

# ========== 第一部分：现状盘点 ==========
add_heading(doc, '一、现状盘点', 1)
add_para(doc, '截至2026年4月，以 OpenClaw 为核心的知识体系已完成基础建设，覆盖数据采集、清洗加工、存储检索、可视化全链路。', size=10)

# 1.1 知识资产
add_heading(doc, '1.1 已有知识资产', 2)
assets = [
    ('TinyDB 知识库', '414 条知识点，覆盖可信数据空间、电力市场、虚拟电厂等核心领域，含来源标注和 CDN 溯源链接'),
    ('知识图谱可视化', '60+ 实体、80+ 关系，实时同步至 GitHub Pages，在线访问'),
    ('长期记忆（MEMORY）', '跨会话精选记忆，沉淀电力现货 API、新能源预测等产品核心信息'),
    ('行业资源库', '12 个数据交易所 + 国家部委网址合集、31 省新能源装机基础数据底座'),
]
t1 = doc.add_table(rows=1, cols=2)
t1.style = 'Table Grid'
add_table_header(t1, ['知识资产', '内容与规模'])
for row_data in assets:
    add_table_data_row(t1, row_data)
doc.add_paragraph()

# 1.2 采集通道
add_heading(doc, '1.2 已有采集通道', 2)
channels = [
    ('北极星电力网', 'Browser 快照', '每日 07:15', '✅ 稳定'),
    ('中国能源新闻网', 'Browser 快照', '每日 07:15', '✅ 稳定'),
    ('RMI 落基山研究所', 'web_fetch', '每日 07:15', '✅ 稳定'),
    ('能见 Eknower', 'web_fetch/Browser', '每日 07:15', '⚠️ 不稳定'),
    ('中电联', 'web_fetch/Browser', '每日 07:15', '✅ 稳定'),
    ('微信公众号（RSS）', 'RSS/Atom', '每日 07:15', '✅ 稳定'),
    ('电力市场情报系统', '5Agent 协作', '定时 cron', '✅ 运行中'),
    ('招投标情报', 'bid-intel-collection', '手动触发', '✅ 运行中'),
]
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
add_table_header(t2, ['采集源', '技术方式', '触发时间', '状态'])
for row_data in channels:
    add_table_data_row(t2, row_data)
doc.add_paragraph()

# 1.3 BCF Skill
add_heading(doc, '1.3 业务专家层 Skill（BCF）', 2)
bcfs = [
    ('energy-news-monitor', '能源新闻采集', '六层完整', '稳定'),
    ('wechat-official-accounts-scan', '公众号 RSS 采集', '六层完整', '稳定'),
    ('power-market-intel', '电力市场情报（5Agent）', '六层完整', '运行中'),
    ('data-market-insight', '数据要素市场情报', '六层完整 v2.1', '运行中'),
    ('pre-sales-workflow', '售前全流程（BD/投标）', '六层完整', '稳定'),
    ('knowledge-management', '知识沉淀统一入口', 'TinyDB 唯一写库', '核心枢纽'),
]
t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
add_table_header(t3, ['Skill 名称', '领域', '架构', '状态'])
for row_data in bcfs:
    add_table_data_row(t3, row_data)
doc.add_paragraph()

doc.add_page_break()

# ========== 第二部分：知识工作方式 ==========
add_heading(doc, '二、知识工作方式（核心架构）', 1)
add_para(doc, '整个知识体系以"采集 → 清洗 → 提炼 → 存储 → 应用 → 反馈"的闭环运转，核心特点如下：', size=10)

flow_items = [
    ('入口唯一', '所有知识写入统一走 learn.py，避免多头录入导致的数据不一致'),
    ('自动化闭环', '采集和情报任务由 cron 驱动，23:00 自动沉淀当日所有 BCF Skill 数据到知识库'),
    ('分层存储', '原始数据（7天清理）→ 提炼知识点（永久）→ 知识库 TinyDB（永久）→ 图谱可视化（GitHub Pages）'),
    ('溯源可查', '每条知识点均标注来源 + CDN 归档链接，可随时回溯原文'),
    ('知识驱动业务', '售前、投标、头条内容等应用层直接从知识库消费，不再重复采集'),
]
for title, desc in flow_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'▶ {title}：')
    set_run_font(run, bold=True, size=10)
    run2 = p.add_run(desc)
    set_run_font(run2, size=10)
doc.add_paragraph()

# 2.1 知识生命周期
add_heading(doc, '2.1 知识生命周期流转', 2)
lifecycle_headers = ['采集', '清洗', '提炼', '存储', '检索', '应用', '反馈']
lifecycle_data = [
    '能源网站\n公众号RSS\n政府官网',
    '去重/过滤\n相关度评级',
    '结构化知识点\n(来源+摘要)',
    'TinyDB\n(唯一写入口)',
    '自然语言查询\n+ 图谱筛选',
    '售前/投标\n日报/头条',
    'success_rate\n记录 → 优化',
]
t4 = doc.add_table(rows=2, cols=7)
t4.style = 'Table Grid'
# 表头行
for i, txt in enumerate(lifecycle_headers):
    c = t4.rows[0].cells[i]
    c.text = txt
    set_cell_bg(c, '0066CC')
    set_run_font(c.paragraphs[0].runs[0], size=9, bold=True, color=(255,255,255))
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# 数据行
for i, txt in enumerate(lifecycle_data):
    c = t4.rows[1].cells[i]
    c.text = txt
    set_run_font(c.paragraphs[0].runs[0], size=8)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 2.2 分层存储
add_heading(doc, '2.2 分层存储架构', 2)
storage = [
    ('L0 原始采集', 'memory/YYYY-MM-DD-*.md', '新闻条目/文章列表', '7天后自动清理'),
    ('L1 加工数据', 'memory/（按日期归档）', '带摘要和相关度的清洗结果', '30天压缩为周报'),
    ('L2 情报报告', 'skills/*/memory/', 'Agent 生成的日报/分析报告', '永久，按月归档'),
    ('L3 提炼知识点', 'knowledge-points/*.json', '结构化知识点', '永久'),
    ('L4 知识库', 'knowledge/db/knowledge-index.json', '去重后的最终知识库', '永久，唯一写入口'),
    ('L5 知识图谱', 'GitHub Pages (data.js)', '实体关系可视化', '自动同步，永久'),
]
t5 = doc.add_table(rows=1, cols=4)
t5.style = 'Table Grid'
add_table_header(t5, ['层级', '存储位置', '数据形态', '保留策略'])
for row_data in storage:
    add_table_data_row(t5, row_data)
doc.add_paragraph()
doc.add_page_break()

# ========== 第三部分：知识分类体系 ==========
add_heading(doc, '三、知识分类体系', 1)

categories = [
    ('📜 政策类', '国家政策（发改委/能源局/数据局）、地方政策（各省电力市场规则）、行业标准（NDI-TR系列）'),
    ('📊 市场类', '电力市场（现货/中长期/辅助服务）、数据交易市场、绿电/绿证市场、储能市场'),
    ('🔧 技术类', '可信数据空间技术、预测算法（电价/新能源/负荷）、AI/大模型、新能源技术（光伏/风电/储能）'),
    ('💼 商业类', '企业画像（朗新/南网/国网）、商业模式（虚拟电厂/充电运营/售电）、竞品分析、招投标情报'),
    ('📈 运营类', '售前材料模板、拜访纪要、会议纪要、头条运营数据'),
]
for cat, desc in categories:
    p = doc.add_paragraph()
    run = p.add_run(f'{cat}')
    set_run_font(run, bold=True, size=11)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(desc)
    set_run_font(run, size=10)
doc.add_paragraph()

# 3.1 入库规则
add_heading(doc, '3.1 知识入库过滤规则', 2)
rules = [
    ('✅ 全部入库', '政策文件（p/miit/ndrc/nea等）、招投标（行业关键词+公告类型）'),
    ('❌ 排除不入库', '新闻/商业媒体（今日头条、网易、新浪财经等噪音内容）'),
]
t6 = doc.add_table(rows=1, cols=2)
t6.style = 'Table Grid'
add_table_header(t6, ['规则', '说明'])
for row_data in rules:
    add_table_data_row(t6, row_data)
doc.add_paragraph()
doc.add_page_break()

# ========== 第四部分：实施路径 ==========
add_heading(doc, '四、实施路径', 1)

phases = [
    {
        'title': 'Phase 1：基础加固（当前 ~ 2026-05-15）',
        'goal': '消除薄弱点，确保所有链路稳定运行',
        'tasks': [
            ('🔴 高', '能见度 Eknower 稳定性修复', '当前标注不稳定，替换采集方案'),
            ('🔴 高', '知识库质量评估机制建立', '414条知识点 success_rate 全为 null，需建立评估机制'),
            ('🟡 中', 'BCF Skill 执行报告规范落地', '每个 Skill 执行后必须报结果到知识管理层'),
            ('🟡 中', 'memory 文件 7 天自动清理机制', '避免原始数据无限积累'),
        ]
    },
    {
        'title': 'Phase 2：能力增强（2026-05-15 ~ 2026-07-01）',
        'goal': '扩展采集源、打通知识应用场景',
        'tasks': [
            ('🔴 高', '新增采集源：各省电力交易中心官网、国家能源局派出机构、储能行业媒体', '扩大情报覆盖范围'),
            ('🔴 高', '知识问答系统', '钉钉提问 → 自动检索知识库 → 带溯源回答'),
            ('🟡 中', '知识点 success_rate 体系', '记录被引用次数、应用成功率'),
            ('🟡 中', '投标场景知识自动匹配', 'pre-sales-workflow 调用知识库，自动匹配政策/案例'),
        ]
    },
    {
        'title': 'Phase 3：体系化运营（2026-07-01 及以后）',
        'goal': '知识体系全面服务于业务，形成自动化知识运营闭环',
        'tasks': [
            ('🔴 高', '知识库 → 飞书/钉钉文档自动同步', '按主题生成多维表格或文档，供团队共享'),
            ('🔴 高', '知识驱动的内容自动生成', '基于知识点自动生成头条文章、周报、月报'),
            ('🟡 中', '客户画像知识自动更新', '拜访前自动从知识库拉取相关政策和竞品情报'),
            ('🟡 中', '知识体系健康度仪表盘', '可视化展示知识点增长、采集覆盖率、质量评分'),
        ]
    },
]

for phase in phases:
    add_heading(doc, phase['title'], 2)
    p = doc.add_paragraph()
    run = p.add_run(f'目标：{phase["goal"]}')
    set_run_font(run, size=10, color=(80, 80, 80))
    t7 = doc.add_table(rows=1, cols=3)
    t7.style = 'Table Grid'
    add_table_header(t7, ['优先级', '任务', '说明'])
    for row_data in phase['tasks']:
        add_table_data_row(t7, row_data)
    doc.add_paragraph()

doc.add_page_break()

# ========== 第五部分：协同开发支撑 ==========
add_heading(doc, '五、协同开发工作流（面向未来）', 1)
add_para(doc, '随着 Trae CN 本地 IDE 与服务器的 SSH 协同链路打通，知识体系将成为"钉钉 → OpenClaw → 本地 Trae"协同开发模式的中枢。', size=10)

workflow_desc = [
    ('需求提出', '小海哥在钉钉提出开发需求'),
    ('知识检索', 'Domi 从 TinyDB 知识库检索相关政策、历史经验、技术方案'),
    ('任务拆解', '将大需求拆解为可执行的子任务（并行 subagent 执行）'),
    ('协同开发', '子任务分流至服务器端（OpenClaw exec）和本地 Trae CN（SSH 直写文件）'),
    ('质量保障', '独立 subagent 做全量测试，通过后再通知用户'),
    ('经验沉淀', '开发完成后自动调用 learn.py 沉淀经验教训到知识库'),
]
for i, (step, desc) in enumerate(workflow_desc, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {step}：')
    set_run_font(run, bold=True, size=11)
    run2 = p.add_run(desc)
    set_run_font(run2, size=11)

doc.add_paragraph()
add_para(doc, '关键支撑点：知识库不再是静态档案，而是贯穿"需求理解 → 方案生成 → 开发执行 → 复盘沉淀"全流程的活性资产。', size=10)
doc.add_page_break()

# ========== 第六部分：核心文件速查 ==========
add_heading(doc, '六、核心文件路径速查', 1)
paths = [
    ('知识库（唯一数据源）', 'knowledge/db/knowledge-index.json'),
    ('知识沉淀入口（唯一写库脚本）', 'knowledge-management/scripts/learn.py'),
    ('知识图谱可视化', 'https://domi-openclaw.github.io/code-space/knowledge-graph/'),
    ('知识图谱同步脚本', 'code-space/knowledge-graph/sync-generator.py'),
    ('长期记忆', 'MEMORY.md'),
    ('日常记录（按日期归档）', 'memory/YYYY-MM-DD.md'),
    ('能源新闻采集', 'memory/YYYY-MM-DD-energy-news.md'),
    ('公众号文章采集', 'memory/YYYY-MM-DD-wechat-articles.md'),
    ('GitHub 文件库（全文归档）', 'Domi-OpenClaw/file-storage'),
    ('GitHub 图床 CDN', 'Domi-OpenClaw/img-bed-picture-storage'),
]
t8 = doc.add_table(rows=1, cols=2)
t8.style = 'Table Grid'
add_table_header(t8, ['组件', '路径/地址'])
for row_data in paths:
    add_table_data_row(t8, row_data)
doc.add_paragraph()

# ========== 结语 ==========
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 本规划立足现有资产，不重复造轮子，所有增强建议均可基于现有架构平滑演进。')
set_run_font(run, size=10, color=(100, 100, 100))

# ========== 保存 ==========
output_path = '/home/admin/.openclaw/workspace/skills/knowledge-management/knowledge/知识体系规划方案_v1.0.docx'
doc.save(output_path)
print(f'✅ 已保存：{output_path}')
